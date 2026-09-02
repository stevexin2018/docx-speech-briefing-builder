#!/usr/bin/env python3
"""
render_docx.py
==============
Word 渲染排版引擎：
将 Markdown 工程技术报告解析并排版为专业优雅的 Word 文档 (.docx)。
包含 11 阶段 LaTeX -> Unicode/OMML 转换、参数表格渲染、大纲标题层级美化与引用高亮块。
已增加 ASCII 边框线与伪表格线过滤，杜绝排版杂音。
"""

import re
import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, border_args in kwargs.items():
        val = border_args.get('val', 'single')
        color = border_args.get('color', 'auto')
        sz = border_args.get('sz', '4')
        b_xml = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders.append(b_xml)
    tcPr.append(tcBorders)

_BS = '\\\\'

def _latex_frac_to_unicode(text):
    for _ in range(20):
        pattern = _BS + r'frac\{([^{}]*)\}\{([^{}]*)\}'
        new_text = re.sub(pattern, r'(\1/\2)', text)
        if new_text == text:
            break
        text = new_text
    return text

def _latex_sqrt_to_unicode(text):
    for _ in range(10):
        pattern = _BS + r'sqrt\{([^{}]*)\}'
        new_text = re.sub(pattern, r'√(\1)', text)
        if new_text == text:
            break
        text = new_text
    return text

def clean_inline_text(text):
    """深度清理行内 LaTeX 标记，转换为规范干净的 Unicode 工程符号，同时严格保护文件路径与行内代码"""
    if not text:
        return ""

    # Phase 0: 保护行内代码与所有系统/文件路径（防止路径反斜杠及文件夹名被 LaTeX 规则误删）
    tokens = {}
    token_idx = [0]

    def _save_token(val):
        key = f"XQZTOKENDOCX{token_idx[0]}QZX"
        tokens[key] = val
        token_idx[0] += 1
        return key

    # 0.1 保护行内代码块 `...`
    text = re.sub(r'`[^`\r\n]+`', lambda m: _save_token(m.group(0)), text)

    # 0.2 保护 Windows 绝对/相对文件路径（如 C:\Users\..., D:/SW2026/... 等）
    win_path_pattern = r'(?<![a-zA-Z0-9_])([a-zA-Z]:[\\/](?:[^\s，。！？；：、`"\'<>\[\]\{\}\(\)\$\^]+[ \\/]?)*[^\s，。！？；：、`"\'<>\[\]\{\}\(\)\$\^]+[\\/]?)'
    text = re.sub(win_path_pattern, lambda m: _save_token(m.group(0)), text)

    # 0.3 保护注册表路径（如 HKLM:\SOFTWARE\..., HKCU:\...）
    reg_pattern = r'(?<![a-zA-Z0-9_])((?:HKLM|HKCU|HKCR|HKU|HKEY_[A-Z_]+):[\\/][^\s，。！？；：、`"\'<>\[\]\{\}\(\)\$\^]+)'
    text = re.sub(reg_pattern, lambda m: _save_token(m.group(0)), text)

    # 0.4 保护 UNC 网络共享路径与 Unix 绝对路径（如 \\server\share, /root/...）
    unc_unix_pattern = r'(?<![a-zA-Z0-9_])(?:\\\\[a-zA-Z0-9_.-]+[\\/]|/(?:root|home|usr|var|etc|opt|tmp|bin|sbin)/)[^\s，。！？；：、`"\'<>\[\]\{\}\(\)\$\^]+'
    text = re.sub(unc_unix_pattern, lambda m: _save_token(m.group(0)), text)

    # Phase 0.5: 剥离数学定界符与温度/角度符号规范化
    text = re.sub(r'\$\$([^$]+)\$\$', r'\1', text)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    text = re.sub(r'[\r\n]?ightarrow', '→', text)
    text = re.sub(r'[\r\n]?ight', '', text)
    # 清除所有空 LaTeX 花括号（如 {}°C -> °C）。
    text = re.sub(r'\{\s*\}', '', text)
    # 温度与度数 LaTeX / 简写结构转换 (如 +5^\circ\text{C}, 5^\circ C, 5^{\circ}\text{C}, 5\degree C -> 5°C)
    text = re.sub(r'\^?\{?' + _BS + r'(?:circ|degree)\}?\s*(?:' + _BS + r'(?:text|mathrm)?\{?([CFcf])\}?|([CFcf]))', r'°\1\2', text)
    text = re.sub(r'\^?\{?' + _BS + r'(?:circ|degree)\}?', '°', text)
    # 修复孤立的 5^ / +5^ 后跟汉字、空格或标点误用 (如 +5^ 环境下 -> +5°C 环境下)
    text = re.sub(r'([+\-]?\d+(?:\.\d+)?)\^(?=[ \t]*[CFcf]\b)', r'\1°', text)
    text = re.sub(r'([+\-]?\d+(?:\.\d+)?)\^(?=[ \t]*[\u4e00-\u9fa5（\(\)，。！？；：`\'\"]|$)', r'\1°C', text)

    # Phase 1: 数学函数名
    math_funcs = [
        'cos', 'sin', 'tan', 'cot', 'sec', 'csc',
        'arccos', 'arcsin', 'arctan',
        'cosh', 'sinh', 'tanh',
        'log', 'ln', 'exp', 'lg',
        'min', 'max', 'sup', 'inf',
        'lim', 'det', 'dim', 'ker', 'deg',
        'gcd', 'lcm', 'mod',
    ]
    for func in sorted(math_funcs, key=len, reverse=True):
        text = re.sub(_BS + func + r'(?![a-zA-Z])', func, text)

    # Phase 2: \text{} 等格式命令
    for _ in range(5):
        new_text = re.sub(_BS + r'(?:text|mathrm|mathbf|mathit|textbf|textit|boldsymbol|operatorname)\{([^{}]*)\}', r'\1', text)
        if new_text == text:
            break
        text = new_text

    # Phase 3: Greek 字母替换
    greek_map = {
        'varepsilon': 'ε', 'varphi': 'φ',
        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'Gamma': 'Γ',
        'delta': 'δ', 'Delta': 'Δ', 'epsilon': 'ε',
        'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'Theta': 'Θ',
        'iota': 'ι', 'kappa': 'κ', 'lambda': 'λ', 'Lambda': 'Λ',
        'mu': 'μ', 'nu': 'ν', 'xi': 'ξ', 'Xi': 'Ξ',
        'pi': 'π', 'Pi': 'Π', 'rho': 'ρ',
        'sigma': 'σ', 'Sigma': 'Σ', 'tau': 'τ', 'upsilon': 'υ',
        'phi': 'φ', 'Phi': 'Φ', 'chi': 'χ', 'psi': 'ψ', 'Psi': 'Ψ',
        'omega': 'ω', 'Omega': 'Ω',
        'infty': '∞', 'partial': '∂', 'nabla': '∇',
    }
    for cmd in sorted(greek_map.keys(), key=len, reverse=True):
        text = re.sub(_BS + cmd + r'(?![a-zA-Z])', greek_map[cmd], text)

    # Phase 4: 上下标花括号简化
    for _ in range(5):
        new_text = re.sub(r'_\{([^{}]+)\}', r'_\1', text)
        new_text = re.sub(r'\^\{([^{}]+)\}', r'^\1', new_text)
        if new_text == text:
            break
        text = new_text

    # Phase 5: 分数与根号
    text = _latex_frac_to_unicode(text)
    text = _latex_sqrt_to_unicode(text)
    text = text.replace('(1/4)', '¼').replace('(1/2)', '½').replace('(3/4)', '¾')

    # Phase 6: 定界符
    text = re.sub(_BS + r'left\s*\(', '(', text)
    text = re.sub(_BS + r'right\s*\)', ')', text)
    text = re.sub(_BS + r'left\s*\[', '[', text)
    text = re.sub(_BS + r'right\s*\]', ']', text)
    text = re.sub(_BS + r'left\s*' + _BS + r'\{', '{', text)
    text = re.sub(_BS + r'right\s*' + _BS + r'\}', '}', text)
    text = re.sub(_BS + r'left', '', text)
    text = re.sub(_BS + r'right', '', text)

    # Phase 7: 运算符号
    symbol_map = {
        r'leq(?![a-zA-Z])': '≤',
        r'le(?![a-zA-Z])': '≤',
        r'geq(?![a-zA-Z])': '≥',
        r'ge(?![a-zA-Z])': '≥',
        r'neq(?![a-zA-Z])': '≠',
        r'ne(?![a-zA-Z])': '≠',
        r'approx(?![a-zA-Z])': '≈',
        r'sim(?![a-zA-Z])': '～',
        r'equiv(?![a-zA-Z])': '≡',
        r'times(?![a-zA-Z])': '×',
        r'cdot(?![a-zA-Z])': '·',
        r'div(?![a-zA-Z])': '÷',
        r'pm(?![a-zA-Z])': '±',
        r'mp(?![a-zA-Z])': '∓',
        r'circ(?![a-zA-Z])': '°',
        r'to(?![a-zA-Z])': '→',
    }
    for cmd, sym in symbol_map.items():
        text = re.sub(_BS + cmd, sym, text)

    # 清理残留
    text = re.sub(r'[\r\n]?ight', '', text)
    text = re.sub(_BS + r'[a-zA-Z]+\{([^{}]*)\}', r'\1', text)
    text = re.sub(_BS + r'[a-zA-Z]+', '', text)
    text = text.replace('\\', '')
    text = text.replace('$', '')
    text = text.replace('^°', '°')
    text = re.sub(r'\^', '', text)
    # 彻底清除重复符号与无意义分割线
    text = re.sub(r'[\-+_=~*#|]{2,}', '', text)

    # Phase 8: 恢复所有受保护的代码块与文件路径令牌
    for key, val in reversed(tokens.items()):
        text = text.replace(key, val)

    return text.strip()


def create_docx_document(topic_title, chapter_id, markdown_content, output_path, theme_color="0F4C81"):
    doc = docx.Document()
    
    # 页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # 页眉
    header = doc.sections[0].header
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_head = p_head.add_run("工程技术标准精读报告 | 专业工程技术资料")
    run_head.font.name = "微软雅黑"
    run_head.font.size = Pt(8.5)
    run_head.font.color.rgb = RGBColor(140, 140, 140)

    # 页脚
    footer = doc.sections[0].footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run("内部受控工程资料 · 请勿外传")
    run_foot.font.name = "微软雅黑"
    run_foot.font.size = Pt(8.5)
    run_foot.font.color.rgb = RGBColor(160, 160, 160)

    # 1. 顶部 Header Banner
    table_head = doc.add_table(rows=1, cols=1)
    table_head.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_h = table_head.cell(0, 0)
    set_cell_background(cell_h, theme_color)
    set_cell_border(cell_h, top={'val':'none'}, bottom={'val':'none'}, left={'val':'none'}, right={'val':'none'})

    p_title = cell_h.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(14)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(topic_title)
    run_title.font.name = "微软雅黑"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)

    p_sub = cell_h.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run(f"【条款/专题定位】{chapter_id}")
    run_sub.font.name = "微软雅黑"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(220, 235, 245)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2. 高级 Markdown 结构解析
    lines = markdown_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        line_s = line.strip()

        # 过滤空行、纯符号分隔线与伪表格对齐线 (+----+----+ / ---------------- / ======)
        if not line_s or line_s == '---' or re.match(r'^[\-+_=\*~#|\s]{2,}$', line_s) or re.match(r'^\|[\s:\-\+\|=]+\|$', line_s):
            i += 1
            continue

        # 独立公式块 $$...$$
        if line_s.startswith('$$'):
            formula_lines = [line_s[2:].strip()]
            if line_s.endswith('$$') and len(line_s) > 4:
                formula_lines = [line_s[2:-2].strip()]
            else:
                i += 1
                while i < len(lines):
                    fl = lines[i].strip()
                    if fl.endswith('$$'):
                        formula_lines.append(fl[:-2].strip())
                        break
                    formula_lines.append(fl)
                    i += 1
            formula_text = ' '.join(f for f in formula_lines if f)
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            set_cell_background(cell, "F0F4F8")
            set_cell_border(cell, left={'val': 'single', 'sz': '24', 'color': theme_color},
                            top={'val': 'none'}, bottom={'val': 'none'}, right={'val': 'none'})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(clean_inline_text(formula_text))
            run.font.name = "Cambria Math"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(30, 30, 30)
            run.font.italic = True
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # 表格识别与渲染 (Markdown Table)
        if line_s.startswith('|') and '|' in line_s[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            rows_data = []
            for tline in table_lines:
                if re.match(r'^\|[\s:\-\+\|=]+\|$', tline):
                    continue
                cells = [c.strip() for c in tline.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)

            if rows_data:
                col_count = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=col_count)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for r_idx, row in enumerate(rows_data):
                    is_header = (r_idx == 0)
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < col_count:
                            cell = table.cell(r_idx, c_idx)
                            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
                            cell.paragraphs[0].paragraph_format.space_after = Pt(4)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                            run = cell.paragraphs[0].add_run(clean_inline_text(cell_text))
                            run.font.name = "微软雅黑"
                            if is_header:
                                run.font.size = Pt(10)
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                set_cell_background(cell, theme_color)
                            else:
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(50, 50, 50)
                                fill = "F0F4F8" if r_idx % 2 == 1 else "FFFFFF"
                                set_cell_background(cell, fill)
                            set_cell_border(cell, top={'val':'single','sz':'4','color':'CCCCCC'},
                                            bottom={'val':'single','sz':'4','color':'CCCCCC'},
                                            left={'val':'single','sz':'4','color':'CCCCCC'},
                                            right={'val':'single','sz':'4','color':'CCCCCC'})
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # 标题 ###
        if line_s.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_inline_text(line_s.replace('### ', '')))
            run.font.name = "微软雅黑"
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 76, 129)
            i += 1
            continue

        # 标题 ####
        if line_s.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_inline_text(line_s.replace('#### ', '')))
            run.font.name = "微软雅黑"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(40, 60, 90)
            i += 1
            continue

        # 引用块 >
        if line_s.startswith('> '):
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            set_cell_background(cell, "F7F9FA")
            set_cell_border(cell, left={'val': 'single', 'sz': '24', 'color': theme_color},
                            top={'val': 'none'}, bottom={'val': 'none'}, right={'val': 'none'})
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_inline_text(line_s.replace('> ', '')))
            run.font.name = "楷体"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25

        cleaned_line = clean_inline_text(line_s)
        parts = cleaned_line.split('**')
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = "微软雅黑"
            run.font.size = Pt(10.5)
            if idx % 2 == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(180, 40, 40)
            else:
                run.font.color.rgb = RGBColor(40, 40, 40)
        i += 1

    doc.save(output_path)
    print(f"Word 文档排版升级构建成功: {output_path}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Render Markdown to Professional Word Document.")
    parser.add_argument("--version", action="version", version="docx-speech-briefing-builder v1.2.7")
    parser.add_argument("--title", default="工程技术报告", help="Title")
    parser.add_argument("--topic-id", default="BPVC", help="Topic ID")
    parser.add_argument("--input", required=True, help="Input markdown path")
    parser.add_argument("--output", required=True, help="Output docx path")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        md_text = f.read()
    create_docx_document(args.title, args.topic_id, md_text, args.output)
